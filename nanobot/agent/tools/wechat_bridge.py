"""WeChat bridge tools for AI-driven operations via MaaWxAuto bridge REST API."""

from typing import Any

from nanobot.agent.tools.base import Tool

_DEFAULT_BRIDGE = "http://127.0.0.1:9574"


class WeChatListSessionsTool(Tool):
    """List visible WeChat chat sessions."""

    def __init__(self, bridge_url: str = _DEFAULT_BRIDGE):
        self._base = bridge_url.rstrip("/")

    @property
    def name(self) -> str:
        return "wechat_list_sessions"

    @property
    def description(self) -> str:
        return (
            "List all visible WeChat chat sessions (contacts and groups). "
            "Returns session names which can be used as chat_id for sending messages."
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {"type": "object", "properties": {}, "required": []}

    async def execute(self, **kwargs: Any) -> str:
        import httpx

        try:
            async with httpx.AsyncClient(timeout=5.0) as client:
                r = await client.get(f"{self._base}/sessions")
                r.raise_for_status()
            sessions = r.json()
            if not sessions:
                return "No WeChat sessions visible."
            lines = [f"- {s['name']} (unread: {s.get('unread', 0)})" for s in sessions]
            return f"WeChat sessions ({len(sessions)}):\n" + "\n".join(lines)
        except Exception as e:
            return f"Error querying bridge: {e}"


class WeChatAddListenerTool(Tool):
    """Add a chat listener to monitor new messages."""

    def __init__(self, bridge_url: str = _DEFAULT_BRIDGE):
        self._base = bridge_url.rstrip("/")

    @property
    def name(self) -> str:
        return "wechat_add_listener"

    @property
    def description(self) -> str:
        return (
            "Start monitoring a WeChat chat for new messages. "
            "The chat will be added to the listen list and new messages "
            "will be forwarded to the bot automatically."
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "chat_id": {
                    "type": "string",
                    "description": "Name of the chat/contact/group to monitor",
                },
            },
            "required": ["chat_id"],
        }

    async def execute(self, chat_id: str, **kwargs: Any) -> str:
        import httpx

        try:
            async with httpx.AsyncClient(timeout=10.0) as client:
                r = await client.post(
                    f"{self._base}/listeners/add",
                    json={"chat_id": chat_id},
                )
                r.raise_for_status()
            data = r.json()
            if data.get("success"):
                return f"Now listening to '{chat_id}': {data.get('message', 'OK')}"
            return f"Failed to add listener: {data.get('message', 'unknown error')}"
        except Exception as e:
            return f"Error: {e}"


class WeChatHealthTool(Tool):
    """Check MaaWxAuto bridge health status."""

    def __init__(self, bridge_url: str = _DEFAULT_BRIDGE):
        self._base = bridge_url.rstrip("/")

    @property
    def name(self) -> str:
        return "wechat_health"

    @property
    def description(self) -> str:
        return (
            "Check the MaaWxAuto bridge health: whether WeChat is connected, "
            "DB is available, and monitoring is active."
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {"type": "object", "properties": {}, "required": []}

    async def execute(self, **kwargs: Any) -> str:
        import httpx

        try:
            async with httpx.AsyncClient(timeout=5.0) as client:
                r = await client.get(f"{self._base}/health")
                r.raise_for_status()
            data = r.json()
            lines = [
                f"Status: {data.get('status', 'unknown')}",
                f"WeChat connected: {data.get('wechat_connected', False)}",
                f"DB available: {data.get('db_available', False)}",
            ]
            return "\n".join(lines)
        except Exception as e:
            return f"Error: Bridge unreachable ({e})"


class ReadFileContentTool(Tool):
    """Read and extract text content from a file (docx, xlsx, pdf, txt, csv)."""

    @property
    def name(self) -> str:
        return "read_file_content"

    @property
    def description(self) -> str:
        return (
            "Read a file from disk and extract its text content. "
            "Supports: .docx, .xlsx, .pdf, .txt, .csv, .md, .json. "
            "Use this to read files shared in WeChat chats when you have the file path."
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Absolute path to the file on disk",
                },
                "max_chars": {
                    "type": "integer",
                    "description": "Maximum characters to return (default 50000)",
                },
            },
            "required": ["file_path"],
        }

    async def execute(self, file_path: str, max_chars: int = 50000, **kw: Any) -> str:
        from pathlib import Path

        p = Path(file_path)
        if not p.exists():
            return f"Error: File not found: {file_path}"
        suffix = p.suffix.lower()
        try:
            if suffix == ".docx":
                return self._read_docx(p, max_chars)
            elif suffix == ".xlsx":
                return self._read_xlsx(p, max_chars)
            elif suffix == ".pdf":
                return self._read_pdf(p, max_chars)
            elif suffix in (".txt", ".csv", ".md", ".json", ".log", ".xml"):
                return p.read_text(encoding="utf-8", errors="replace")[:max_chars]
            else:
                return f"Error: Unsupported file type: {suffix}"
        except Exception as e:
            return f"Error reading {p.name}: {e}"

    @staticmethod
    def _read_docx(p: Any, max_chars: int) -> str:
        import docx

        doc = docx.Document(str(p))
        lines = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        text = "\n".join(lines)
        return text[:max_chars]

    @staticmethod
    def _read_xlsx(p: Any, max_chars: int) -> str:
        import openpyxl

        wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in wb.worksheets:
            parts.append(f"=== Sheet: {sheet.title} ===")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                parts.append("\t".join(cells))
        wb.close()
        text = "\n".join(parts)
        return text[:max_chars]

    @staticmethod
    def _read_pdf(p: Any, max_chars: int) -> str:
        import pdfplumber

        text_parts: list[str] = []
        with pdfplumber.open(str(p)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                for table in page.extract_tables():
                    for row in table:
                        cells = [str(c) if c else "" for c in row]
                        text_parts.append(" | ".join(cells))
        text = "\n".join(text_parts)
        return text[:max_chars]


class ExportPdfTool(Tool):
    """Export Markdown content as a styled PDF file (Chinese-friendly)."""

    @property
    def name(self) -> str:
        return "export_pdf"

    @property
    def description(self) -> str:
        return (
            "将 Markdown 内容导出为 PDF 文件。用于 WeChat 长回复场景——"
            "当回复超过 500 字或包含表格/分析时，先导出 PDF 再通过 message tool 的 media 参数发送文件。"
            "返回生成的 PDF 文件绝对路径。"
        )

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "content": {
                    "type": "string",
                    "description": "Markdown formatted text to export as PDF",
                },
                "title": {
                    "type": "string",
                    "description": "Filename prefix (default: 'report'). Chinese OK.",
                },
            },
            "required": ["content"],
        }

    async def execute(self, content: str, title: str = "report", **kw: Any) -> str:
        import tempfile
        from datetime import datetime
        from pathlib import Path

        import markdown as md
        from fpdf import FPDF

        # --- helpers --------------------------------------------------------
        def _find_cjk_font() -> tuple[Path | None, str]:
            """Return (font_path, family_name) for the first available CJK font."""
            candidates = [
                (Path("C:/Windows/Fonts/msyh.ttc"), "msyh"),
                (Path("C:/Windows/Fonts/simhei.ttf"), "simhei"),
                (Path("C:/Windows/Fonts/simsun.ttc"), "simsun"),
            ]
            for p, family in candidates:
                if p.exists():
                    return p, family
            return None, ""

        # --- convert MD → HTML ---------------------------------------------
        html = md.markdown(
            content,
            extensions=["tables", "fenced_code", "nl2br", "sane_lists"],
        )

        # --- build PDF ------------------------------------------------------
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        font_path, font_family = _find_cjk_font()
        if font_path:
            pdf.add_font(font_family, "", str(font_path), uni=True)
            pdf.add_font(font_family, "B", str(font_path), uni=True)
            pdf.add_font(font_family, "I", str(font_path), uni=True)
            pdf.set_font(font_family, size=11)
        else:
            pdf.set_font("Helvetica", size=11)

        pdf.write_html(html)

        # --- save -----------------------------------------------------------
        export_dir = Path(tempfile.gettempdir()) / "nanobot_exports"
        export_dir.mkdir(parents=True, exist_ok=True)

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Sanitise title for filesystem
        safe_title = "".join(c if c.isalnum() or c in "._-" else "_" for c in title)
        filename = f"{safe_title}_{ts}.pdf"
        out_path = export_dir / filename
        pdf.output(str(out_path))

        return f"PDF exported: {out_path}"


# All WeChat bridge tool classes for easy registration
WECHAT_BRIDGE_TOOLS = [
    WeChatListSessionsTool,
    WeChatAddListenerTool,
    WeChatHealthTool,
]

# Tools that don't require bridge_url (registered separately)
WECHAT_LOCAL_TOOLS = [
    ReadFileContentTool,
    ExportPdfTool,
]
